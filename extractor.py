import os
import pytesseract
from PIL import Image
import fitz  # PyMuPDF
from pdf2image import convert_from_path
from docx import Document
import pandas as pd
import zipfile
import tempfile
import shutil

class UniversalTextExtractor:
    """
    Complete Universal Text Extractor for all supported file types.
    Handles PDF, DOCX, Excel, Images, and ZIP files with robust error handling.
    """

    def __init__(self):
        self.supported_formats = {
            'images': ['.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.gif'],
            'pdfs': ['.pdf'],
            'word': ['.docx', '.doc'],
            'excel': ['.xlsx', '.xls', '.csv'],
            'archives': ['.zip']
        }
        print("🚀 UniversalTextExtractor initialized")
        print(f"📁 Supported formats: {sum(len(formats) for formats in self.supported_formats.values())} types")

    def extract_text(self, file_path):
        """
        Universal text extractor for all supported file types.
        Returns extracted text or error message.
        """
        if not os.path.exists(file_path):
            return f"Error: File '{file_path}' not found"

        file_ext = os.path.splitext(file_path.lower())[1]
        print(f"📄 Processing file: {os.path.basename(file_path)}")
        print(f"📋 Detected format: {file_ext}")

        try:
            # Determine file type and extract accordingly
            if file_ext in self.supported_formats['images']:
                return self._extract_from_image(file_path)
            elif file_ext in self.supported_formats['pdfs']:
                return self._extract_from_pdf(file_path)
            elif file_ext in self.supported_formats['word']:
                return self._extract_from_word(file_path)
            elif file_ext in self.supported_formats['excel']:
                return self._extract_from_excel(file_path)
            elif file_ext in self.supported_formats['archives']:
                return self._extract_from_zip(file_path)
            else:
                return f"Error: Unsupported file format '{file_ext}'"

        except Exception as e:
            error_msg = f"Error processing {os.path.basename(file_path)}: {str(e)}"
            print(f"❌ {error_msg}")
            return error_msg

    def _extract_from_image(self, image_path):
        """Extract text from image using OCR"""
        try:
            print("🖼️  Using OCR for image...")
            image = Image.open(image_path)

            # Try multiple OCR configurations for better accuracy
            configs = [
                '--oem 3 --psm 6',  # Default
                '--oem 3 --psm 3',  # Fully automatic page segmentation
                '--oem 3 --psm 1',  # Automatic page segmentation with OSD
            ]

            best_text = ""
            for config in configs:
                try:
                    text = pytesseract.image_to_string(image, lang='eng', config=config)
                    if len(text.strip()) > len(best_text.strip()):
                        best_text = text
                except:
                    continue

            # Clean up the text
            cleaned_text = '\n'.join(line.strip() for line in best_text.split('\n') if line.strip())
            result = cleaned_text if cleaned_text else "No text found in image"
            print(f"✅ OCR extraction complete: {len(cleaned_text)} characters")
            return result

        except Exception as e:
            return f"OCR Error: {str(e)}"

    def _extract_from_pdf(self, pdf_path):
        """Extract text from PDF with fallback to OCR"""
        try:
            print("📄 Processing PDF...")

            # METHOD 1: Try direct text extraction first
            print("📝 Trying direct text extraction...")
            doc = fitz.open(pdf_path)
            full_text = []
            text_found = False

            for page_num, page in enumerate(doc, 1):
                text = page.get_text().strip()
                if text:
                    text_found = True
                    full_text.append(text)
                    print(f"✅ Found text on page {page_num} ({len(text)} chars)")

            doc.close()

            if text_found:
                result = "\n\n".join(full_text)
                print(f"✅ Direct extraction successful: {len(result)} characters")
                return result

            # METHOD 2: If no text found, use OCR
            print("🔍 No direct text found. Using OCR...")
            try:
                pages = convert_from_path(pdf_path, dpi=200, first_page=1, last_page=5)  # Limit pages for demo
                ocr_texts = []

                for page_num, page in enumerate(pages, 1):
                    print(f"🖼️  OCR processing page {page_num}/{len(pages)}...")

                    # Multiple OCR attempts for better accuracy
                    configs = ['--oem 3 --psm 6', '--oem 3 --psm 3']
                    best_page_text = ""

                    for config in configs:
                        try:
                            text = pytesseract.image_to_string(page, lang='eng', config=config)
                            if len(text.strip()) > len(best_page_text.strip()):
                                best_page_text = text
                        except:
                            continue

                    cleaned_text = '\n'.join(line.strip() for line in best_page_text.split('\n') if line.strip())
                    if cleaned_text:
                        ocr_texts.append(cleaned_text)

                if ocr_texts:
                    result = "\n\n".join(ocr_texts)
                    print(f"✅ OCR extraction successful: {len(result)} characters")
                    return result
                else:
                    return "No text could be extracted from this PDF"

            except Exception as ocr_error:
                return f"PDF OCR failed: {str(ocr_error)}"

        except Exception as e:
            return f"Error processing PDF: {str(e)}"

    def _extract_from_word(self, word_path):
        """Extract text from Word documents"""
        try:
            print("📝 Processing Word document...")
            document = Document(word_path)

            # Extract paragraphs
            paragraphs = []
            for para in document.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Extract tables
            table_data = []
            for table in document.tables:
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_data.append(cell.text.strip())
                    if row_data:
                        table_data.append(' | '.join(row_data))

            # Extract headers and footers
            header_footer_text = []
            for section in document.sections:
                # Headers
                if section.header:
                    for para in section.header.paragraphs:
                        if para.text.strip():
                            header_footer_text.append(f"[HEADER] {para.text}")

                # Footers
                if section.footer:
                    for para in section.footer.paragraphs:
                        if para.text.strip():
                            header_footer_text.append(f"[FOOTER] {para.text}")

            # Combine all text
            all_text = paragraphs + table_data + header_footer_text
            full_text = '\n'.join(all_text)

            print(f"✅ Word extraction complete: {len(paragraphs)} paragraphs, {len(document.tables)} tables")
            return full_text if full_text else "No text found in Word document"

        except Exception as e:
            return f"Error processing Word document: {str(e)}"

    def _extract_from_excel(self, excel_path):
        """Extract text from Excel files"""
        try:
            print("📊 Processing Excel file...")

            # Handle CSV files
            if excel_path.lower().endswith('.csv'):
                df = pd.read_csv(excel_path)
                excel_data = {'Sheet1': df}
            else:
                # Load all sheets
                excel_data = pd.read_excel(excel_path, sheet_name=None)

            all_text = []
            total_rows = 0

            for sheet_name, df in excel_data.items():
                print(f"📋 Processing sheet: {sheet_name}")

                # Add sheet header
                all_text.append(f"=== SHEET: {sheet_name} ===")

                # Include column headers
                headers = ' | '.join(str(col) for col in df.columns if str(col) != 'nan')
                all_text.append(f"[HEADERS] {headers}")

                # Extract all rows
                for index, row in df.iterrows():
                    row_values = [str(cell) for cell in row.values if pd.notna(cell) and str(cell).strip()]
                    if row_values:
                        row_text = ' | '.join(row_values)
                        all_text.append(row_text)
                        total_rows += 1

                all_text.append("")  # Empty line between sheets

            full_text = '\n'.join(all_text)
            print(f"✅ Excel extraction complete: {len(excel_data)} sheets, {total_rows} rows")
            return full_text if full_text else "No text found in Excel file"

        except Exception as e:
            return f"Error processing Excel file: {str(e)}"

    def _extract_from_zip(self, zip_path):
        """Extract and process files from ZIP archive"""
        try:
            print("🗂️  Processing ZIP archive...")

            with tempfile.TemporaryDirectory() as temp_dir:
                # Extract ZIP contents
                with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                    zip_ref.extractall(temp_dir)

                # Get all extracted files
                extracted_files = []
                for root, dirs, files in os.walk(temp_dir):
                    for file in files:
                        file_path = os.path.join(root, file)
                        file_ext = os.path.splitext(file.lower())[1]

                        # Only process supported file types
                        all_supported = []
                        for format_list in self.supported_formats.values():
                            all_supported.extend(format_list)

                        if file_ext in all_supported and file_ext != '.zip':  # Avoid nested ZIP
                            extracted_files.append(file_path)

                print(f"📁 Found {len(extracted_files)} processable files in ZIP")

                # Process each file
                all_extracted_text = []
                processed_count = 0

                for file_path in extracted_files[:10]:  # Limit to first 10 files for demo
                    try:
                        relative_path = os.path.relpath(file_path, temp_dir)
                        print(f"📄 Processing: {relative_path}")

                        # Recursively extract text (but avoid ZIP files)
                        file_ext = os.path.splitext(file_path.lower())[1]
                        if file_ext != '.zip':
                            text = self.extract_text(file_path)

                            if not text.startswith("Error:") and text.strip():
                                all_extracted_text.append(f"=== FILE: {relative_path} ===")
                                all_extracted_text.append(text)
                                all_extracted_text.append("")  # Empty line
                                processed_count += 1

                    except Exception as file_error:
                        print(f"⚠️  Failed to process {relative_path}: {file_error}")
                        continue

                if all_extracted_text:
                    result = '\n'.join(all_extracted_text)
                    print(f"✅ ZIP extraction complete: {processed_count} files processed")
                    return result
                else:
                    return "No text could be extracted from files in ZIP archive"

        except Exception as e:
            return f"Error processing ZIP file: {str(e)}"

    def get_supported_extensions(self):
        """Get list of all supported file extensions"""
        all_extensions = []
        for format_list in self.supported_formats.values():
            all_extensions.extend(format_list)
        return sorted(all_extensions)

    def is_supported(self, file_path):
        """Check if file type is supported"""
        file_ext = os.path.splitext(file_path.lower())[1]
        return file_ext in self.get_supported_extensions()

# Standalone function for backward compatibility
def extract_text_from_file(file_path: str) -> str:
    """
    Simple wrapper function for text extraction.
    Used by other modules that need text extraction.
    """
    extractor = UniversalTextExtractor()
    return extractor.extract_text(file_path)

# Demo and testing
if __name__ == "__main__":
    print("🚀 UNIVERSAL TEXT EXTRACTOR - DEMO")
    print("=" * 60)

    # Configuration
    TEST_FILE = "sample1.pdf"  # Change this to test different files

    # Initialize extractor
    extractor = UniversalTextExtractor()

    print(f"🎯 Testing with file: {TEST_FILE}")
    print(f"📋 Supported extensions: {', '.join(extractor.get_supported_extensions())}")

    # Check if file exists and is supported
    if os.path.exists(TEST_FILE):
        if extractor.is_supported(TEST_FILE):
            print(f"✅ File is supported")

            # Extract text
            print("\n" + "=" * 60)
            print("EXTRACTION RESULTS:")
            print("=" * 60)

            extracted_text = extractor.extract_text(TEST_FILE)

            if not extracted_text.startswith("Error:"):
                print(f"✅ Extraction successful!")
                print(f"📊 Text length: {len(extracted_text)} characters")
                print(f"📝 Lines count: {len(extracted_text.splitlines())}")
                print("\n" + "-" * 40)
                print("PREVIEW (first 500 characters):")
                print("-" * 40)
                print(extracted_text[:500] + "..." if len(extracted_text) > 500 else extracted_text)
            else:
                print(f"❌ Extraction failed: {extracted_text}")
        else:
            print(f"❌ File type not supported: {os.path.splitext(TEST_FILE)[1]}")
    else:
        print(f"❌ Test file not found: {TEST_FILE}")
        print("\n🔧 Available test files:")

        # List available files in current directory
        supported_files = []
        for file in os.listdir('.'):
            if extractor.is_supported(file):
                supported_files.append(file)

        if supported_files:
            for file in supported_files[:5]:  # Show first 5
                print(f"   📄 {file}")
        else:
            print("   No supported files found in current directory")

    print("\n" + "=" * 60)
    print("🎉 DEMO COMPLETE")
    print("=" * 60)